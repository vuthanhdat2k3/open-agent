from __future__ import annotations

from app.customer_intelligence.contracts import ReportSections


def render_markdown(sections: ReportSections) -> str:
    """Render the stable 7-section briefing as canonical markdown (FR-006).

    Every external claim in the output is traceable to ``sections["sources"]``;
    absent sections render as explicit empty markers rather than fabrications.
    """
    out: list[str] = []
    out.append("# Customer Intelligence Briefing")
    out.append("")

    out.append("## Executive Summary")
    out.append(sections.get("executive_summary") or "_No summary available._")
    out.append("")

    out.append("## Company Overview")
    overview = sections.get("company_overview") or []
    if overview:
        for company in overview:
            out.append(f"### {company.get('canonical_name') or 'Unknown company'}")
            if company.get("aliases"):
                out.append(f"- Aliases: {', '.join(company['aliases'])}")
            if company.get("industry"):
                out.append(f"- Industry: {company['industry']}")
            if company.get("products"):
                out.append(f"- Products: {', '.join(company['products'])}")
            if company.get("domain"):
                out.append(f"- Domain: {company['domain']}")
    else:
        out.append("_No company matched._")
    out.append("")

    out.append("## Recent News")
    news = sections.get("recent_news") or []
    if news:
        for item in news:
            date = f" ({item.get('published_date')})" if item.get("published_date") else ""
            out.append(f"- [{item.get('title') or item.get('url')}]({item.get('url')}){date}")
            if item.get("excerpt"):
                out.append(f"  {item['excerpt'][:300]}")
    else:
        out.append("_No recent news found._")
    out.append("")

    out.append("## Contact Information")
    contacts = sections.get("contact_information") or []
    if contacts:
        for contact in contacts:
            name = contact.get("name") or "Unknown"
            role = f" — {contact.get('role')}" if contact.get("role") else ""
            email = f" ({contact.get('email')})" if contact.get("email") else ""
            out.append(f"- {name}{role}{email}")
    else:
        out.append("_No contacts found._")
    out.append("")

    out.append("## Upcoming Meetings")
    meetings = sections.get("upcoming_meetings") or []
    if meetings:
        for meeting in meetings:
            when = meeting.get("start_at") or "unscheduled"
            match = f" [{meeting.get('match_type')}]" if meeting.get("match_type") else ""
            out.append(f"- {meeting.get('title') or '(untitled)'} — {when}{match}")
            if meeting.get("attendees"):
                out.append(f"  Attendees: {', '.join(meeting['attendees'])}")
    else:
        out.append("_No upcoming meetings matched._")
    out.append("")

    out.append("## Open Questions")
    questions = sections.get("open_questions") or []
    if questions:
        for question in questions:
            out.append(f"- {question}")
    else:
        out.append("_None._")
    out.append("")

    out.append("## Sources")
    sources = sections.get("sources") or []
    if sources:
        for i, source in enumerate(sources, start=1):
            title = source.get("title") or source.get("url")
            date = f" ({source.get('published_date')})" if source.get("published_date") else ""
            publisher = f" — {source.get('publisher')}" if source.get("publisher") else ""
            out.append(f"{i}. [{title}]({source.get('url')}){date}{publisher}")
            if source.get("excerpt"):
                out.append(f"   {source['excerpt'][:300]}")
    else:
        out.append("_No sources._")

    return "\n".join(out)


# --------------------------------------------------------------------------- #
# On-demand artifact renderers
# --------------------------------------------------------------------------- #

import html as _html
from io import BytesIO


def _escape(value: object) -> str:
    return _html.escape(str(value), quote=True)


def render_html(sections: ReportSections) -> str:
    """Render the seven report sections as a self-contained, deterministic HTML document."""
    out = [
        "<!doctype html>",
        '<html lang="en"><head><meta charset="utf-8">',
        "<title>Customer Intelligence Briefing</title>",
        "<style>body{font-family:Arial,sans-serif;color:#18202a;line-height:1.5;max-width:900px;margin:2rem auto;padding:0 1rem}h1{color:#123c69}h2{border-bottom:1px solid #d8dee8;padding-bottom:.3rem;margin-top:2rem}h3{margin-bottom:.2rem}li{margin:.25rem 0}.source{margin:.5rem 0}.muted{color:#667085}</style>",
        "</head><body>",
        "<h1>Customer Intelligence Briefing</h1>",
        "<h2>Executive Summary</h2>",
        f"<p>{_escape(sections.get('executive_summary') or 'No summary available.')}</p>",
    ]

    out.append("<h2>Company Overview</h2>")
    overview = sections.get("company_overview") or []
    if overview:
        for company in overview:
            out.append(f"<h3>{_escape(company.get('canonical_name') or 'Unknown company')}</h3><ul>")
            for label, key in (("Aliases", "aliases"), ("Industry", "industry"), ("Products", "products"), ("Domain", "domain")):
                value = company.get(key)
                if value:
                    rendered = ", ".join(str(item) for item in value) if isinstance(value, list) else str(value)
                    out.append(f"<li><strong>{label}:</strong> {_escape(rendered)}</li>")
            out.append("</ul>")
    else:
        out.append('<p class="muted">No company matched.</p>')

    out.append("<h2>Recent News</h2>")
    news = sections.get("recent_news") or []
    if news:
        out.append("<ul>")
        for item in news:
            title = item.get("title") or item.get("url") or "Untitled source"
            url = item.get("url") or "#"
            date = f" ({_escape(item['published_date'])})" if item.get("published_date") else ""
            out.append(f'<li><a href="{_escape(url)}">{_escape(title)}</a>{date}')
            if item.get("excerpt"):
                out.append(f"<br><span class=\"muted\">{_escape(str(item['excerpt'])[:300])}</span>")
            out.append("</li>")
        out.append("</ul>")
    else:
        out.append('<p class="muted">No recent news found.</p>')

    out.append("<h2>Contact Information</h2>")
    contacts = sections.get("contact_information") or []
    if contacts:
        out.append("<ul>")
        for contact in contacts:
            value = _escape(contact.get("name") or "Unknown")
            if contact.get("role"):
                value += f" — {_escape(contact['role'])}"
            if contact.get("email"):
                value += f" ({_escape(contact['email'])})"
            out.append(f"<li>{value}</li>")
        out.append("</ul>")
    else:
        out.append('<p class="muted">No contacts found.</p>')

    out.append("<h2>Upcoming Meetings</h2>")
    meetings = sections.get("upcoming_meetings") or []
    if meetings:
        out.append("<ul>")
        for meeting in meetings:
            value = _escape(meeting.get("title") or "(untitled)")
            value += f" — {_escape(meeting.get('start_at') or 'unscheduled')}"
            if meeting.get("match_type"):
                value += f" [{_escape(meeting['match_type'])}]"
            out.append(f"<li>{value}")
            if meeting.get("attendees"):
                out.append(f"<br><span class=\"muted\">Attendees: {_escape(', '.join(meeting['attendees']))}</span>")
            out.append("</li>")
        out.append("</ul>")
    else:
        out.append('<p class="muted">No upcoming meetings matched.</p>')

    out.append("<h2>Open Questions</h2>")
    questions = sections.get("open_questions") or []
    if questions:
        out.append("<ul>" + "".join(f"<li>{_escape(question)}</li>" for question in questions) + "</ul>")
    else:
        out.append('<p class="muted">None.</p>')

    out.append("<h2>Sources</h2>")
    sources = sections.get("sources") or []
    if sources:
        out.append("<ol>")
        for source in sources:
            title = source.get("title") or source.get("url") or "Untitled source"
            url = source.get("url") or "#"
            date = f" ({_escape(source['published_date'])})" if source.get("published_date") else ""
            publisher = f" — {_escape(source['publisher'])}" if source.get("publisher") else ""
            out.append(f'<li class="source"><a href="{_escape(url)}">{_escape(title)}</a>{date}{publisher}')
            if source.get("excerpt"):
                out.append(f"<br><span class=\"muted\">{_escape(str(source['excerpt'])[:300])}</span>")
            out.append("</li>")
        out.append("</ol>")
    else:
        out.append('<p class="muted">No sources.</p>')

    out.append("</body></html>")
    return "".join(out)


def render_pdf(html: str) -> bytes:
    """Convert self-contained report HTML to PDF using pure-Python xhtml2pdf."""
    try:
        from xhtml2pdf import pisa
    except ImportError as exc:  # pragma: no cover - dependency is pinned in pyproject
        raise RuntimeError("PDF rendering requires xhtml2pdf") from exc

    output = BytesIO()
    result = pisa.CreatePDF(html, dest=output, encoding="utf-8")
    if result.err:
        raise ValueError("could not render briefing PDF")
    return output.getvalue()


def render_docx(sections: ReportSections) -> bytes:
    """Build a DOCX directly from structured report sections."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency is pinned in pyproject
        raise RuntimeError("DOCX rendering requires python-docx") from exc

    document = Document()
    document.add_heading("Customer Intelligence Briefing", level=0)
    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(sections.get("executive_summary") or "No summary available.")

    document.add_heading("Company Overview", level=1)
    overview = sections.get("company_overview") or []
    if overview:
        for company in overview:
            document.add_heading(str(company.get("canonical_name") or "Unknown company"), level=2)
            for label, key in (("Aliases", "aliases"), ("Industry", "industry"), ("Products", "products"), ("Domain", "domain")):
                value = company.get(key)
                if value:
                    rendered = ", ".join(str(item) for item in value) if isinstance(value, list) else str(value)
                    document.add_paragraph(f"{label}: {rendered}", style="List Bullet")
    else:
        document.add_paragraph("No company matched.")

    document.add_heading("Recent News", level=1)
    for item in sections.get("recent_news") or []:
        title = str(item.get("title") or item.get("url") or "Untitled source")
        date = f" ({item['published_date']})" if item.get("published_date") else ""
        document.add_paragraph(f"{title}{date}", style="List Bullet")
        if item.get("excerpt"):
            document.add_paragraph(str(item["excerpt"])[:300])
    if not sections.get("recent_news"):
        document.add_paragraph("No recent news found.")

    document.add_heading("Contact Information", level=1)
    contacts = sections.get("contact_information") or []
    for contact in contacts:
        value = str(contact.get("name") or "Unknown")
        if contact.get("role"):
            value += f" — {contact['role']}"
        if contact.get("email"):
            value += f" ({contact['email']})"
        document.add_paragraph(value, style="List Bullet")
    if not contacts:
        document.add_paragraph("No contacts found.")

    document.add_heading("Upcoming Meetings", level=1)
    meetings = sections.get("upcoming_meetings") or []
    for meeting in meetings:
        value = f"{meeting.get('title') or '(untitled)'} — {meeting.get('start_at') or 'unscheduled'}"
        if meeting.get("match_type"):
            value += f" [{meeting['match_type']}]"
        document.add_paragraph(value, style="List Bullet")
        if meeting.get("attendees"):
            document.add_paragraph(f"Attendees: {', '.join(meeting['attendees'])}")
    if not meetings:
        document.add_paragraph("No upcoming meetings matched.")

    document.add_heading("Open Questions", level=1)
    questions = sections.get("open_questions") or []
    for question in questions:
        document.add_paragraph(str(question), style="List Bullet")
    if not questions:
        document.add_paragraph("None.")

    document.add_heading("Sources", level=1)
    sources = sections.get("sources") or []
    for source in sources:
        title = str(source.get("title") or source.get("url") or "Untitled source")
        publisher = f" — {source['publisher']}" if source.get("publisher") else ""
        date = f" ({source['published_date']})" if source.get("published_date") else ""
        document.add_paragraph(f"{title}{date}{publisher}", style="List Number")
        if source.get("url"):
            document.add_paragraph(str(source["url"]))
    if not sources:
        document.add_paragraph("No sources.")

    output = BytesIO()
    document.save(output)
    return output.getvalue()
