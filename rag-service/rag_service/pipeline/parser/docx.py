"""DOCX parser.

Iterates paragraphs (preserving heading levels as markdown) and renders tables
as markdown pipe tables. Uses ``python-docx`` lazily.
"""

from __future__ import annotations

from rag_service.core.logging import logger
from rag_service.exceptions import ParseError
from rag_service.pipeline.base import Parser, ParseResult

__all__ = ["DOCXParser"]


def _heading_level(style_name: str) -> int:
    name = (style_name or "").lower()
    if name.startswith("heading"):
        digits = "".join(ch for ch in name if ch.isdigit())
        if digits:
            return int(digits)
        return 1
    if "title" in name:
        return 1
    return 0


class DOCXParser(Parser):
    """Parse DOCX ``bytes`` (and ``.docm``) into markdown-flavoured text."""

    async def parse(self, source: bytes | str, **kwargs: object) -> ParseResult:
        if isinstance(source, str):
            source = source.encode("utf-8", errors="ignore")

        try:
            import io

            import docx  # type: ignore
        except ImportError as exc:
            raise ParseError(
                "DOCX parsing requires 'python-docx' to be installed"
            ) from exc

        try:
            document = docx.Document(io.BytesIO(source))
        except Exception as exc:
            raise ParseError(f"Failed to read DOCX document: {exc}") from exc

        parts: list[str] = []
        word_count = 0

        for para in document.paragraphs:
            text = (para.text or "").strip()
            style_name = ""
            try:
                style_name = (para.style and para.style.name) or ""
            except Exception:  # pragma: no cover - defensive
                style_name = ""
            level = _heading_level(style_name)
            if level:
                parts.append(f"{'#' * min(level, 6)} {text}" if text else "")
            else:
                parts.append(text)
            if text:
                word_count += len(text.split())

        for table in document.tables:
            parts.append("")
            parts.append(_table_to_markdown(table))
            parts.append("")

        body = "\n".join(parts).strip()

        metadata: dict[str, object] = {}
        core = getattr(document, "core_properties", None)
        if core is not None:
            title = getattr(core, "title", None)
            if title:
                metadata["title"] = str(title)
            author = getattr(core, "author", None)
            if author:
                metadata["author"] = str(author)
            created = getattr(core, "created", None)
            if created is not None:
                metadata["created"] = created.isoformat() if hasattr(created, "isoformat") else str(created)
            modified = getattr(core, "modified", None)
            if modified is not None:
                metadata["modified"] = modified.isoformat() if hasattr(modified, "isoformat") else str(modified)
        metadata["word_count"] = word_count

        return ParseResult(text=body, metadata=metadata)


def _table_to_markdown(table: object) -> str:
    rows: list[list[str]] = []
    try:
        for row in table.rows:
            cells = [("".join(c.text for c in cell.paragraphs)).strip() for cell in row.cells]
            rows.append(cells)
    except Exception as exc:  # pragma: no cover - defensive
        logger.warning("Failed to read a DOCX table: %s", exc)
        return ""

    if not rows:
        return ""

    out: list[str] = []
    header = rows[0]
    out.append("| " + " | ".join(header) + " |")
    out.append("| " + " | ".join("---" for _ in header) + " |")
    for row in rows[1:]:
        out.append("| " + " | ".join(row) + " |")
    return "\n".join(out)
